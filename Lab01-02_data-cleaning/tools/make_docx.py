# -*- coding: utf-8 -*-
"""สร้างไฟล์รายงาน Lab1 & Lab2 เป็น .docx"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = "TH Sarabun New"      # ฟอนต์ราชการไทย (ถ้าไม่มีจะ fallback อัตโนมัติ)
FALLBACK = "Tahoma"

doc = Document()

# ---- ตั้งค่าฟอนต์ default ให้รองรับภาษาไทย ----
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(15)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), FONT)
rfonts.set(qn("w:hAnsi"), FONT)
rfonts.set(qn("w:cs"), FONT)      # complex script (ไทย)


def set_run_font(run, size=15, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:cs"), FONT)


def add_heading(text, size, color=(31, 78, 121), space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    return p


def add_para(text, size=15, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_bullet(text, size=15):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p


def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, size=14, bold=True, color=(255, 255, 255))
    # shade header
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:fill"): "1F4E79"})
        tcPr.append(shd)
    # body
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            set_run_font(r, size=14)
    return table


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("รายงาน Lab 1 & Lab 2 — Analytics Engineering")
set_run_font(r, size=20, bold=True, color=(31, 78, 121))

# =====================================================================
# LAB 1
# =====================================================================
add_heading("Lab 1: จำแนกประเภทและออกแบบการจัดเก็บข้อมูล", 17)
add_para("ตารางวิเคราะห์ข้อมูล 8 ประเภทของมหาวิทยาลัย")

lab1_headers = ["#", "ข้อมูล", "ประเภทข้อมูล", "รูปแบบไฟล์ / Schema (อธิบาย)", "แหล่งจัดเก็บที่เหมาะสม"]
lab1_rows = [
    ["1", "ตารางข้อมูลนิสิต (CSV)", "Structured",
     "ไฟล์ CSV มี schema ชัดเจน คอลัมน์คงที่ เช่น student_id, name, faculty, gpa แต่ละแถวคือ 1 record",
     "Relational Database (MySQL/PostgreSQL) หรือ Data Warehouse เมื่อใช้วิเคราะห์"],
    ["2", "ข้อมูลเซนเซอร์อุณหภูมิ (JSON)", "Semi-structured",
     "JSON แบบ key–value / nested เช่น {device_id, temp_c, ts} schema ยืดหยุ่น เปลี่ยน field ได้",
     "Data Lake หรือ NoSQL (MongoDB) รองรับข้อมูลไหลเข้าต่อเนื่อง"],
    ["3", "ภาพจากกล้องวงจรปิด", "Unstructured",
     "ไฟล์ภาพ/วิดีโอ (JPG, MP4) ไม่มี schema เป็น binary ขนาดใหญ่",
     "Object Storage (Amazon S3, MinIO, Azure Blob)"],
    ["4", "ไฟล์ PDF ใบสมัครทุน", "Unstructured",
     "เอกสาร PDF ข้อความ+ภาพ ไม่มีโครงสร้างตาราง (ต้องใช้ OCR/NLP ดึงข้อมูล)",
     "Object Storage หรือ Data Lake"],
    ["5", "Log การเข้าใช้งาน LMS", "Semi-structured",
     "ไฟล์ log กึ่งมีรูปแบบ เช่น timestamp user_id action ip หรือ JSON log line",
     "Data Lake (เก็บดิบ) → Data Warehouse เมื่อทำ analytics"],
    ["6", "ผลการลงทะเบียนเรียน (MySQL)", "Structured",
     "ตารางในฐานข้อมูลเชิงสัมพันธ์ schema ชัด มี PK/FK เช่น enrollment(student_id, course_id, grade)",
     "Relational Database (MySQL) / Data Warehouse สำหรับรายงาน"],
    ["7", "ไฟล์เสียงศูนย์รับแจ้งปัญหา", "Unstructured",
     "ไฟล์เสียง (WAV, MP3) เป็น binary ไม่มี schema (ต้องแปลงเป็นข้อความด้วย Speech-to-Text)",
     "Object Storage"],
    ["8", "ข้อมูลดิบเครื่องอ่านบัตร (ยังไม่ตรวจสอบ)", "Raw Data (มัก Semi-structured)",
     "ข้อมูลดิบยังไม่ถูก clean/validate อาจมี error, ค่าว่าง, รูปแบบไม่สม่ำเสมอ",
     "Data Lake (Raw / Landing Zone) เก็บต้นฉบับก่อนทำความสะอาด"],
]
add_table(lab1_headers, lab1_rows)

add_para("สรุปหลักการเลือกแหล่งจัดเก็บ", bold=True, space_after=2)
add_bullet("Structured → Relational Database / Data Warehouse")
add_bullet("Semi-structured → Data Lake / NoSQL")
add_bullet("Unstructured → Object Storage")
add_bullet("Raw Data → Data Lake (Raw Zone) ก่อนแปลงเป็น Cleaned/Curated Zone")

# =====================================================================
# LAB 2
# =====================================================================
add_heading("Lab 2: ทำความสะอาดข้อมูลดิบจากเซนเซอร์ IoT", 17, space_before=18)
add_para("ไฟล์ต้นฉบับ: iot_sensor_raw_data_extended.csv (480 แถวข้อมูล + 1 header)")

# 1) ปัญหา
add_heading("1) ปัญหาคุณภาพข้อมูลที่พบ (รายงานปัญหา)", 15)
prob_headers = ["ประเภทปัญหา", "ตัวอย่างที่พบในไฟล์", "รายละเอียด"]
prob_rows = [
    ["Error – เซนเซอร์เสีย", "temp_c = -999.0", "ค่า sentinel ที่เซนเซอร์ส่งเมื่ออ่านค่าไม่ได้"],
    ["Error – motion ผิดรูปแบบ", "FALSEE, TRUEE, unknown", "สะกดผิด / ค่าที่ไม่ใช่ true-false"],
    ["Outlier – temp_c นอกช่วง", "85.0, 67.8", "เกินช่วง 0–50 °C ที่กำหนด"],
    ["Outlier – battery นอกช่วง", "-5, -1, 108, 115, 999", "เกินช่วง 0–100"],
    ["Duplicate", "2026-07-12 14:53:00, SN-9982 ปรากฏ 2 ครั้ง", "แถวซ้ำที่มี device_id + timestamp เดียวกัน (25 คู่)"],
    ["Missing Value", "NULL ใน device_id, temp_c, motion, battery", "ค่าว่างในหลายคอลัมน์"],
]
add_table(prob_headers, prob_rows)
add_para("หมายเหตุ: ค่า motion ที่แปลงได้ เช่น yes→true, 0→false, 1→true ถือเป็นการ normalize ไม่ใช่ error", size=13)

# 2) Rules
add_heading("2) Data Quality Rules ที่กำหนด", 15)
add_para("1. temp_c ต้องอยู่ระหว่าง 0–50 °C (ค่า -999 = sensor error, นอกช่วง = outlier → NULL)")
add_para("2. device_id ห้ามเป็นค่าว่าง (ถ้าว่าง = แถว INVALID ระบุแหล่งที่มาไม่ได้)")
add_para("3. ห้ามมีข้อมูลซ้ำที่ device_id + timestamp เดียวกัน")
add_para("4. battery ต้องอยู่ระหว่าง 0–100 (นอกช่วง → NULL)")
add_para("5. motion ต้องเป็น true / false เท่านั้น (ค่าผิดรูป → NULL)")

# 3) โปรแกรม
add_heading("3) โปรแกรมทำความสะอาด (Python / Pandas)", 15)
add_para("Source code: clean_iot_data.py ทำงาน 5 ขั้นตอน")
add_bullet("แปลง token ค่าว่าง (NULL, \"\", NaN) → NaN จริง")
add_bullet("Normalize motion เป็น true/false, ค่าผิดรูป → NULL")
add_bullet("แปลง temp_c / battery เป็นตัวเลข, ค่า error/outlier → NULL")
add_bullet("ลบข้อมูลซ้ำ (device_id + timestamp)")
add_bullet("เพิ่มคอลัมน์ data_quality_status (OK / MISSING_VALUE / INVALID_NO_DEVICE_ID)")
add_para("การจัดการ Missing Value: เลือกทำเครื่องหมายสถานะ (flag) ผ่านคอลัมน์ data_quality_status "
         "แทนการเติมค่ามั่ว เพื่อไม่บิดเบือนข้อมูลเซนเซอร์ ผู้ใช้ปลายทางเลือกได้เองว่าจะกรองแถว "
         "MISSING_VALUE ออก หรือเติมค่า (เช่น interpolate ตามเวลา)", size=13)

# 4) เปรียบเทียบ
add_heading("4) เปรียบเทียบจำนวนแถวก่อน–หลังทำความสะอาด", 15)
add_table(["รายการ", "จำนวนแถว"], [
    ["ก่อนทำความสะอาด", "480"],
    ["ลบข้อมูลซ้ำออก", "−25"],
    ["หลังทำความสะอาด", "455"],
])

# 5) สรุปปัญหา
add_heading("5) ตารางสรุปจำนวนปัญหา", 15)
add_table(["ประเภทปัญหา", "จำนวน"], [
    ["Error (ค่าผิด/เซนเซอร์เสีย: -999 + motion ผิดรูป)", "9"],
    ["Duplicate (ข้อมูลซ้ำ)", "25"],
    ["Missing Value (ค่าว่าง หลังแปลง error/outlier)", "32"],
    ["Outlier (temp/battery นอกช่วง)", "9"],
])
add_para("การกระจายค่า data_quality_status หลังทำความสะอาด:", bold=True, space_after=2)
add_table(["สถานะ", "จำนวนแถว"], [
    ["OK", "423"],
    ["MISSING_VALUE", "30"],
    ["INVALID_NO_DEVICE_ID", "2"],
])

# 6) เหตุผลไม่ลบ raw
add_heading("6) เหตุผลที่ไม่ควรลบ Raw Data ต้นฉบับ", 15)
add_para("1. ตรวจสอบย้อนกลับได้ (Traceability/Audit): เก็บต้นฉบับเพื่อพิสูจน์ที่มาและตรวจสอบภายหลัง")
add_para("2. แก้ไข Logic การ clean ใหม่ได้: ถ้ากฎคุณภาพเปลี่ยน ต้องกลับไปประมวลผลจาก Raw ใหม่")
add_para("3. ข้อมูลที่ดูเหมือนผิดอาจมีความหมาย: เช่น -999 บอกช่วงที่เซนเซอร์เสีย สำคัญเชิงบำรุงรักษา")
add_para("4. ป้องกันข้อมูลสูญหายถาวร: การ clean คือการตีความ ถ้าลบต้นฉบับแล้วตีความผิดจะกู้คืนไม่ได้")
add_para("5. หลักการ Data Lake: แยกชั้น Raw → Cleaned → Curated ต้นฉบับต้องคงอยู่เสมอ")

# ไฟล์ผลลัพธ์
add_heading("ไฟล์ผลลัพธ์ที่ส่ง", 15, space_before=14)
add_table(["ไฟล์", "คำอธิบาย"], [
    ["clean_iot_data.py", "Source code (Python/Pandas)"],
    ["iot_sensor_cleaned_data.csv", "ไฟล์ Cleaned Data (455 แถว + คอลัมน์ data_quality_status)"],
    ["รายงาน_Lab1_Lab2.docx", "รายงานปัญหา + ตารางสรุป + ตอบ Lab 1"],
])

OUT = "รายงาน_Lab1_Lab2.docx"
doc.save(OUT)
print("บันทึกไฟล์แล้ว ->", OUT)
